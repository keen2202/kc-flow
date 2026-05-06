"""文档处理器技能入口"""

from typing import Dict, Any, List
from pydantic import BaseModel, Field


class ProcessInput(BaseModel):
    file_url: str
    file_type: str
    extraction_mode: str = "full_text"
    language: str = "zh-CN"
    max_pages: int = 100


class SectionInfo(BaseModel):
    title: str
    content: str
    page: int


class DocumentMetadata(BaseModel):
    page_count: int
    word_count: int
    author: str = ""
    created_at: str = ""


class ProcessOutput(BaseModel):
    content: str
    metadata: DocumentMetadata
    sections: List[SectionInfo] = []


async def process(input: ProcessInput, context: "SkillContext") -> ProcessOutput:
    """
    文档处理技能入口。

    支持 PDF、DOCX、XLSX、TXT、Markdown 格式的文档解析。
    """
    context.logger.info(f"Processing {input.file_type} document from {input.file_url}")

    # 下载文档
    raw_bytes = await context.http_request("GET", input.file_url)

    # 根据文件类型选择解析器
    match input.file_type:
        case "pdf":
            content, metadata, sections = await _parse_pdf(raw_bytes, input)
        case "docx":
            content, metadata, sections = await _parse_docx(raw_bytes, input)
        case "xlsx":
            content, metadata, sections = await _parse_xlsx(raw_bytes, input)
        case "txt" | "markdown":
            content, metadata, sections = await _parse_text(raw_bytes, input)
        case _:
            raise ValueError(f"Unsupported file type: {input.file_type}")

    return ProcessOutput(
        content=content,
        metadata=metadata,
        sections=sections,
    )


async def _parse_pdf(raw: bytes, input: ProcessInput):
    """解析PDF文档"""
    import io
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    pages = reader.pages[:input.max_pages]

    text_parts = []
    for i, page in enumerate(pages):
        page_text = page.extract_text()
        text_parts.append(page_text)

    content = "\n\n".join(text_parts)
    metadata = DocumentMetadata(
        page_count=len(reader.pages),
        word_count=len(content.split()),
        author=reader.metadata.get("/Author", ""),
        created_at=reader.metadata.get("/CreationDate", ""),
    )

    return content, metadata, []


async def _parse_docx(raw: bytes, input: ProcessInput):
    """解析Word文档"""
    import io
    from docx import Document

    doc = Document(io.BytesIO(raw))
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(text_parts)

    sections = []
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith("Heading"):
            sections.append(SectionInfo(
                title=para.text,
                content="",
                page=i // 40 + 1,
            ))

    metadata = DocumentMetadata(
        page_count=0,
        word_count=len(content.split()),
        author=doc.core_properties.author or "",
        created_at=str(doc.core_properties.created) if doc.core_properties.created else "",
    )

    return content, metadata, sections


async def _parse_xlsx(raw: bytes, input: ProcessInput):
    """解析Excel文档"""
    import io
    import json
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(raw), data_only=True)
    all_data = {}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_data = []
        for row in ws.iter_rows(values_only=True):
            sheet_data.append(list(row))
        all_data[sheet_name] = sheet_data

    content = json.dumps(all_data, ensure_ascii=False, default=str)
    metadata = DocumentMetadata(
        page_count=len(wb.sheetnames),
        word_count=len(str(all_data).split()),
    )

    return content, metadata, []


async def _parse_text(raw: bytes, input: ProcessInput):
    """解析纯文本/Markdown文档"""
    content = raw.decode("utf-8")

    if input.language == "zh-CN":
        word_count = len(content.replace("\n", "").replace(" ", ""))
    else:
        word_count = len(content.split())

    metadata = DocumentMetadata(
        page_count=1,
        word_count=word_count,
    )

    return content, metadata, []
