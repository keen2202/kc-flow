"""Data processing and integration nodes: Code, HTTPRequest, Template, DataTransform, DocumentParser, Webhook, MCPTool, Wait."""

import asyncio
import json
import time
from typing import Any

import structlog

from src.engine.abstractions import (
    BaseNode,
    ExecutionContext,
    NodeCategory,
    NodeResult,
    NodeStatus,
    VariableDef,
    register_node,
)

logger = structlog.get_logger()


# ──────────────────────────────────────────────
# Code Node
# ──────────────────────────────────────────────


@register_node(
    node_type="code",
    display_name="代码执行",
    category=NodeCategory.DATA,
    icon="code",
    description="在安全沙箱中执行 Python 或 JavaScript 代码",
    inputs=[
        VariableDef(name="code", type="string", required=True, description="要执行的代码"),
        VariableDef(name="inputs", type="object", description="代码输入变量"),
    ],
    outputs=[
        VariableDef(name="result", type="any", description="代码执行返回值"),
        VariableDef(name="stdout", type="string", description="标准输出"),
        VariableDef(name="stderr", type="string", description="标准错误"),
    ],
    config_schema={
        "type": "object",
        "properties": {
            "language": {"type": "string", "enum": ["python", "javascript"], "default": "python"},
            "code": {"type": "string", "description": "代码内容"},
            "dependencies": {"type": "array", "items": {"type": "string"}, "description": "依赖包列表"},
            "timeout": {"type": "integer", "default": 60, "description": "超时时间(秒)"},
        },
        "required": ["language"],
    },
)
class CodeNode(BaseNode):
    """Executes code in a Docker sandbox."""

    async def execute(self, variable_pool: Any, context: ExecutionContext | None = None) -> NodeResult:
        from src.services.sandbox import DockerSandbox, SandboxConfig

        language = self.node_config.get("language", "python")
        code_template = self.node_config.get("code", "")
        timeout = self.node_config.get("timeout", 60)

        # Get code from config or input
        if code_template:
            code = variable_pool.resolve_template(code_template) if "{{" in code_template else code_template
        else:
            code = variable_pool.get(f"{self.node_id}.input.code", "")

        if not code:
            return NodeResult(status=NodeStatus.FAILED, error="No code provided")

        # Get input variables
        inputs = variable_pool.get(f"{self.node_id}.input.inputs", {})

        # Wrap code to capture return value
        if language == "python":
            wrapped_code = self._wrap_python_code(code, inputs)
        else:
            wrapped_code = self._wrap_javascript_code(code, inputs)

        config = SandboxConfig(timeout_seconds=timeout)
        sandbox = DockerSandbox(config)

        try:
            result = await sandbox.execute(code=wrapped_code, language=language)

            if result.timed_out:
                return NodeResult(status=NodeStatus.TIMEOUT, error=f"Code execution timed out after {timeout}s")

            if result.exit_code != 0:
                return NodeResult(
                    status=NodeStatus.FAILED,
                    error=f"Code execution failed (exit {result.exit_code}): {result.stderr}",
                )

            # Parse output
            try:
                output = json.loads(result.stdout)
                return_val = output.get("result")
                stdout = output.get("stdout", "")
            except (json.JSONDecodeError, TypeError):
                return_val = result.stdout
                stdout = result.stdout

            return NodeResult(
                status=NodeStatus.SUCCEEDED,
                outputs={"result": return_val, "stdout": stdout, "stderr": result.stderr},
            )
        except Exception as e:
            logger.error("Code execution failed", node_id=self.node_id, error=str(e))
            return NodeResult(status=NodeStatus.FAILED, error=str(e))

    def _wrap_python_code(self, code: str, inputs: dict) -> str:
        """Wrap Python code to inject inputs and capture output."""
        import_section = "import json, sys\n"
        input_section = f"_inputs = {json.dumps(inputs)}\n"
        # If code has a main() function, call it
        if "def main(" in code:
            call_section = "\n_result = main(**_inputs)\n"
        else:
            call_section = ""
        output_section = """
_stdout_capture = []
class _CaptureStdout:
    def write(self, s):
        _stdout_capture.append(s)
    def flush(self): pass
_old_stdout = sys.stdout
sys.stdout = _CaptureStdout()
"""
        end_section = """
sys.stdout = _old_stdout
print(json.dumps({"result": _result if '_result' in dir() else None, "stdout": "".join(_stdout_capture)}))
"""
        return import_section + input_section + output_section + code + call_section + end_section

    def _wrap_javascript_code(self, code: str, inputs: dict) -> str:
        """Wrap JavaScript code to inject inputs and capture output."""
        return f"""
const _inputs = {json.dumps(inputs)};
let _result;
const _origLog = console.log;
const _stdout = [];
console.log = (...args) => _stdout.push(args.join(' '));
try {{
    {code}
    if (typeof main === 'function') _result = main(_inputs);
}} catch(e) {{
    console.error(e.message);
}}
console.log = _origLog;
process.stdout.write(JSON.stringify({{result: _result, stdout: _stdout.join('\\n')}}));
"""


# ──────────────────────────────────────────────
# HTTP Request Node
# ──────────────────────────────────────────────


@register_node(
    node_type="http_request",
    display_name="HTTP请求",
    category=NodeCategory.DATA,
    icon="globe",
    description="发送 HTTP 请求调用外部 API",
    inputs=[
        VariableDef(name="url", type="string", required=True, description="请求 URL"),
        VariableDef(name="body", type="object", description="请求体"),
        VariableDef(name="headers", type="object", description="请求头"),
    ],
    outputs=[
        VariableDef(name="status_code", type="number", description="HTTP 状态码"),
        VariableDef(name="body", type="any", description="响应体"),
        VariableDef(name="headers", type="object", description="响应头"),
    ],
    config_schema={
        "type": "object",
        "properties": {
            "url": {"type": "string", "description": "请求 URL，支持模板"},
            "method": {"type": "string", "enum": ["GET", "POST", "PUT", "PATCH", "DELETE"], "default": "GET"},
            "headers": {"type": "object", "description": "请求头"},
            "body_template": {"type": "string", "description": "请求体模板"},
            "timeout": {"type": "integer", "default": 30},
            "retry_config": {
                "type": "object",
                "properties": {
                    "max_retries": {"type": "integer", "default": 3},
                    "backoff_factor": {"type": "number", "default": 2.0},
                },
            },
        },
        "required": ["url", "method"],
    },
)
class HTTPRequestNode(BaseNode):
    """Sends HTTP requests to external APIs."""

    @property
    def supported_retry_exceptions(self) -> tuple[type[Exception], ...]:
        return (TimeoutError, ConnectionError, OSError)

    async def execute(self, variable_pool: Any, context: ExecutionContext | None = None) -> NodeResult:
        import httpx

        url_template = self.node_config.get("url", "")
        method = self.node_config.get("method", "GET").upper()
        headers_config = self.node_config.get("headers", {})
        body_template = self.node_config.get("body_template", "")
        timeout = self.node_config.get("timeout", 30)

        # Resolve URL template
        url = variable_pool.resolve_template(url_template) if "{{" in url_template else url_template
        if not url:
            url = variable_pool.get(f"{self.node_id}.input.url", "")

        if not url:
            return NodeResult(status=NodeStatus.FAILED, error="No URL provided")

        # Resolve headers
        headers = {k: variable_pool.resolve_template(str(v)) if "{{" in str(v) else v for k, v in headers_config.items()}

        # Resolve body
        body = None
        if body_template:
            body_str = variable_pool.resolve_template(body_template) if "{{" in body_template else body_template
            try:
                body = json.loads(body_str)
            except (json.JSONDecodeError, TypeError):
                body = body_str
        else:
            body = variable_pool.get(f"{self.node_id}.input.body")

        try:
            async with httpx.AsyncClient(timeout=timeout) as client:
                response = await client.request(
                    method=method,
                    url=url,
                    headers=headers,
                    json=body if isinstance(body, (dict, list)) else None,
                    content=body if isinstance(body, str) else None,
                )

                # Parse response body
                try:
                    response_body = response.json()
                except (json.JSONDecodeError, TypeError):
                    response_body = response.text

                return NodeResult(
                    status=NodeStatus.SUCCEEDED,
                    outputs={
                        "status_code": response.status_code,
                        "body": response_body,
                        "headers": dict(response.headers),
                    },
                )
        except Exception as e:
            logger.error("HTTP request failed", node_id=self.node_id, error=str(e))
            return NodeResult(status=NodeStatus.FAILED, error=str(e))


# ──────────────────────────────────────────────
# Template Node
# ──────────────────────────────────────────────


@register_node(
    node_type="template",
    display_name="模板转换",
    category=NodeCategory.DATA,
    icon="file-text",
    description="使用 Jinja2 模板进行文本转换",
    inputs=[
        VariableDef(name="template", type="string", required=True, description="Jinja2 模板"),
        VariableDef(name="inputs", type="object", description="模板输入变量"),
    ],
    outputs=[
        VariableDef(name="text", type="string", description="渲染后的文本"),
    ],
    config_schema={
        "type": "object",
        "properties": {
            "template": {"type": "string", "description": "Jinja2 模板内容"},
            "template_file": {"type": "string", "description": "模板文件路径"},
            "inputs": {"type": "object", "description": "输入变量"},
            "output_format": {"type": "string", "enum": ["text", "markdown", "html", "json"], "default": "text"},
        },
    },
)
class TemplateNode(BaseNode):
    """Renders Jinja2 templates with variable substitution."""

    async def execute(self, variable_pool: Any, context: ExecutionContext | None = None) -> NodeResult:
        template_str = self.node_config.get("template", "")
        output_format = self.node_config.get("output_format", "text")

        if not template_str:
            template_str = variable_pool.get(f"{self.node_id}.input.template", "")

        if not template_str:
            return NodeResult(status=NodeStatus.FAILED, error="No template provided")

        try:
            rendered = variable_pool.resolve_template(template_str)

            # Post-process based on output format
            if output_format == "json":
                try:
                    rendered = json.loads(rendered)
                except json.JSONDecodeError:
                    pass

            return NodeResult(
                status=NodeStatus.SUCCEEDED,
                outputs={"text": rendered},
            )
        except Exception as e:
            logger.error("Template rendering failed", node_id=self.node_id, error=str(e))
            return NodeResult(status=NodeStatus.FAILED, error=str(e))


# ──────────────────────────────────────────────
# Data Transform Node
# ──────────────────────────────────────────────


@register_node(
    node_type="data_transform",
    display_name="数据转换",
    category=NodeCategory.DATA,
    icon="shuffle",
    description="数据格式转换、映射和聚合",
    inputs=[
        VariableDef(name="data", type="any", required=True, description="输入数据"),
    ],
    outputs=[
        VariableDef(name="result", type="any", description="转换后的数据"),
    ],
    config_schema={
        "type": "object",
        "properties": {
            "input_format": {"type": "string", "enum": ["json", "csv", "text", "xml"], "default": "json"},
            "output_format": {"type": "string", "enum": ["json", "csv", "text", "markdown"], "default": "json"},
            "mapping_rules": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "source": {"type": "string", "description": "源字段路径"},
                        "target": {"type": "string", "description": "目标字段路径"},
                        "transform": {"type": "string", "description": "转换表达式"},
                    },
                    "required": ["source", "target"],
                },
            },
            "filter_expression": {"type": "string", "description": "数据过滤表达式"},
            "aggregation": {
                "type": "object",
                "properties": {
                    "group_by": {"type": "string"},
                    "function": {"type": "string", "enum": ["count", "sum", "avg", "min", "max"]},
                    "field": {"type": "string"},
                },
            },
        },
    },
)
class DataTransformNode(BaseNode):
    """Transforms data between formats with mapping and aggregation."""

    async def execute(self, variable_pool: Any, context: ExecutionContext | None = None) -> NodeResult:
        data = variable_pool.get(f"{self.node_id}.input.data")
        if data is None:
            return NodeResult(status=NodeStatus.FAILED, error="No data provided")

        mapping_rules = self.node_config.get("mapping_rules", [])
        filter_expr = self.node_config.get("filter_expression")
        aggregation = self.node_config.get("aggregation")

        try:
            result = data

            # Apply mapping rules
            if mapping_rules:
                result = self._apply_mapping(result, mapping_rules)

            # Apply filter
            if filter_expr and isinstance(result, list):
                result = self._apply_filter(result, filter_expr)

            # Apply aggregation
            if aggregation and isinstance(result, list):
                result = self._apply_aggregation(result, aggregation)

            return NodeResult(
                status=NodeStatus.SUCCEEDED,
                outputs={"result": result},
            )
        except Exception as e:
            logger.error("Data transform failed", node_id=self.node_id, error=str(e))
            return NodeResult(status=NodeStatus.FAILED, error=str(e))

    def _apply_mapping(self, data: Any, rules: list[dict]) -> Any:
        """Apply field mapping rules to data."""
        if isinstance(data, dict):
            result: dict[str, Any] = {}
            for rule in rules:
                source = rule["source"]
                target = rule["target"]
                value = self._get_nested(data, source)
                if value is not _MISSING:
                    self._set_nested(result, target, value)
            return result
        elif isinstance(data, list):
            return [self._apply_mapping(item, rules) for item in data]
        return data

    def _apply_filter(self, data: list, expression: str) -> list:
        """Filter data items based on expression."""
        resolved = expression
        result = []
        for item in data:
            try:
                if isinstance(item, dict):
                    check = resolved
                    for k, v in item.items():
                        check = check.replace(f"{{{{{k}}}}}", repr(v))
                    allowed_names = {"True": True, "False": False, "None": None}
                    if eval(check, {"__builtins__": {}}, allowed_names):
                        result.append(item)
                else:
                    result.append(item)
            except Exception:
                result.append(item)
        return result

    def _apply_aggregation(self, data: list, aggregation: dict) -> Any:
        """Apply aggregation function to data."""
        group_by = aggregation.get("group_by")
        function = aggregation.get("function", "count")
        field = aggregation.get("field")

        if group_by:
            groups: dict[str, list] = {}
            for item in data:
                if isinstance(item, dict):
                    key = str(item.get(group_by, ""))
                    groups.setdefault(key, []).append(item)

            result = {}
            for key, items in groups.items():
                if function == "count":
                    result[key] = len(items)
                elif field:
                    values = [item.get(field) for item in items if isinstance(item, dict) and field in item]
                    numeric = [v for v in values if isinstance(v, (int, float))]
                    if function == "sum":
                        result[key] = sum(numeric)
                    elif function == "avg":
                        result[key] = sum(numeric) / len(numeric) if numeric else 0
                    elif function == "min":
                        result[key] = min(numeric) if numeric else None
                    elif function == "max":
                        result[key] = max(numeric) if numeric else None
            return result
        else:
            if function == "count":
                return len(data)
            elif field:
                values = [item.get(field) for item in data if isinstance(item, dict) and field in item]
                numeric = [v for v in values if isinstance(v, (int, float))]
                if function == "sum":
                    return sum(numeric)
                elif function == "avg":
                    return sum(numeric) / len(numeric) if numeric else 0
                elif function == "min":
                    return min(numeric) if numeric else None
                elif function == "max":
                    return max(numeric) if numeric else None
            return data

    def _get_nested(self, obj: dict, path: str) -> Any:
        """Get value from nested dict by dot-separated path."""
        keys = path.split(".")
        current = obj
        for key in keys:
            if isinstance(current, dict) and key in current:
                current = current[key]
            else:
                return _MISSING
        return current

    def _set_nested(self, obj: dict, path: str, value: Any) -> None:
        """Set value in nested dict by dot-separated path."""
        keys = path.split(".")
        current = obj
        for key in keys[:-1]:
            if key not in current or not isinstance(current[key], dict):
                current[key] = {}
            current = current[key]
        current[keys[-1]] = value


_MISSING = object()


# ──────────────────────────────────────────────
# Document Parser Node
# ──────────────────────────────────────────────


@register_node(
    node_type="document_parser",
    display_name="文档解析",
    category=NodeCategory.DATA,
    icon="file",
    description="解析 PDF、DOCX、XLSX 等文档格式",
    inputs=[
        VariableDef(name="file_source", type="string", required=True, description="文件路径或URL"),
    ],
    outputs=[
        VariableDef(name="content", type="string", description="提取的文本内容"),
        VariableDef(name="metadata", type="object", description="文档元数据"),
        VariableDef(name="pages", type="array", description="分页内容"),
    ],
    config_schema={
        "type": "object",
        "properties": {
            "file_source": {"type": "string", "description": "文件路径或URL"},
            "file_type": {"type": "string", "enum": ["pdf", "docx", "xlsx", "csv", "txt"], "description": "文件类型"},
            "extraction_mode": {"type": "string", "enum": ["text", "tables", "all"], "default": "text"},
            "page_range": {"type": "string", "description": "页码范围，如 1-10"},
        },
    },
)
class DocumentParserNode(BaseNode):
    """Parses documents (PDF, DOCX, XLSX) and extracts text content."""

    async def execute(self, variable_pool: Any, context: ExecutionContext | None = None) -> NodeResult:
        file_source = self.node_config.get("file_source", "")
        if not file_source:
            file_source = variable_pool.get(f"{self.node_id}.input.file_source", "")

        if not file_source:
            return NodeResult(status=NodeStatus.FAILED, error="No file source provided")

        file_type = self.node_config.get("file_type", "")
        extraction_mode = self.node_config.get("extraction_mode", "text")

        # Auto-detect file type from extension
        if not file_type:
            if "." in file_source:
                file_type = file_source.rsplit(".", 1)[-1].lower()
            else:
                return NodeResult(status=NodeStatus.FAILED, error="Cannot determine file type")

        try:
            if file_type == "pdf":
                return self._parse_pdf(file_source, extraction_mode)
            elif file_type == "docx":
                return self._parse_docx(file_source)
            elif file_type == "xlsx":
                return self._parse_xlsx(file_source)
            elif file_type == "csv":
                return self._parse_csv(file_source)
            elif file_type == "txt":
                return self._parse_text(file_source)
            else:
                return NodeResult(status=NodeStatus.FAILED, error=f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error("Document parsing failed", node_id=self.node_id, error=str(e))
            return NodeResult(status=NodeStatus.FAILED, error=str(e))

    def _parse_pdf(self, file_source: str, mode: str) -> NodeResult:
        from pypdf2 import PdfReader
        reader = PdfReader(file_source)
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
        content = "\n\n".join(pages)
        metadata = {
            "page_count": len(reader.pages),
            "metadata": reader.metadata if reader.metadata else {},
        }
        return NodeResult(
            status=NodeStatus.SUCCEEDED,
            outputs={"content": content, "metadata": metadata, "pages": pages},
        )

    def _parse_docx(self, file_source: str) -> NodeResult:
        from docx import Document
        doc = Document(file_source)
        paragraphs = [p.text for p in doc.paragraphs]
        content = "\n".join(paragraphs)
        metadata = {
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
        }
        return NodeResult(
            status=NodeStatus.SUCCEEDED,
            outputs={"content": content, "metadata": metadata, "pages": [content]},
        )

    def _parse_xlsx(self, file_source: str) -> NodeResult:
        from openpyxl import load_workbook
        wb = load_workbook(file_source, read_only=True)
        sheets_data: list[dict[str, Any]] = []
        all_text: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows: list[list[Any]] = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(cell) if cell is not None else "" for cell in row])
                all_text.append("\t".join(str(cell) if cell is not None else "" for cell in row))
            sheets_data.append({"name": sheet_name, "rows": rows})

        content = "\n".join(all_text)
        metadata = {"sheet_count": len(sheets_data), "sheets": [s["name"] for s in sheets_data]}
        return NodeResult(
            status=NodeStatus.SUCCEEDED,
            outputs={"content": content, "metadata": metadata, "pages": [content]},
        )

    def _parse_csv(self, file_source: str) -> NodeResult:
        import csv
        rows: list[list[str]] = []
        with open(file_source, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(row)
        content = "\n".join(",".join(row) for row in rows)
        metadata = {"row_count": len(rows), "column_count": len(rows[0]) if rows else 0}
        return NodeResult(
            status=NodeStatus.SUCCEEDED,
            outputs={"content": content, "metadata": metadata, "pages": [content]},
        )

    def _parse_text(self, file_source: str) -> NodeResult:
        with open(file_source, "r", encoding="utf-8") as f:
            content = f.read()
        metadata = {"char_count": len(content), "line_count": content.count("\n") + 1}
        return NodeResult(
            status=NodeStatus.SUCCEEDED,
            outputs={"content": content, "metadata": metadata, "pages": [content]},
        )


# ──────────────────────────────────────────────
# Webhook Node
# ──────────────────────────────────────────────


@register_node(
    node_type="webhook",
    display_name="Webhook",
    category=NodeCategory.DATA,
    icon="link",
    description="接收外部 Webhook 回调数据",
    inputs=[
        VariableDef(name="payload", type="object", description="Webhook 接收的 payload"),
    ],
    outputs=[
        VariableDef(name="payload", type="object", description="接收到的数据"),
        VariableDef(name="headers", type="object", description="请求头"),
        VariableDef(name="method", type="string", description="HTTP 方法"),
    ],
    config_schema={
        "type": "object",
        "properties": {
            "webhook_path": {"type": "string", "description": "Webhook 路径"},
            "secret": {"type": "string", "description": "签名验证密钥"},
            "allowed_methods": {
                "type": "array",
                "items": {"type": "string", "enum": ["GET", "POST", "PUT", "DELETE"]},
                "default": ["POST"],
            },
        },
    },
)
class WebhookNode(BaseNode):
    """Receives external webhook callbacks. Data is injected by the API layer."""

    async def execute(self, variable_pool: Any, context: ExecutionContext | None = None) -> NodeResult:
        # Webhook data is injected into the variable pool by the API endpoint
        payload = variable_pool.get(f"{self.node_id}.input.payload", {})
        headers = variable_pool.get(f"{self.node_id}.input.headers", {})
        method = variable_pool.get(f"{self.node_id}.input.method", "POST")

        # Verify secret if configured
        secret = self.node_config.get("secret")
        if secret:
            signature = headers.get("x-webhook-signature", "")
            if not self._verify_signature(payload, secret, signature):
                return NodeResult(status=NodeStatus.FAILED, error="Webhook signature verification failed")

        return NodeResult(
            status=NodeStatus.SUCCEEDED,
            outputs={"payload": payload, "headers": headers, "method": method},
        )

    def _verify_signature(self, payload: Any, secret: str, signature: str) -> bool:
        """Verify webhook signature (HMAC-SHA256)."""
        import hashlib
        import hmac
        expected = hmac.new(
            secret.encode(),
            json.dumps(payload, sort_keys=True).encode(),
            hashlib.sha256,
        ).hexdigest()
        return hmac.compare_digest(expected, signature)


# ──────────────────────────────────────────────
# MCP Tool Node
# ──────────────────────────────────────────────


@register_node(
    node_type="mcp_tool",
    display_name="MCP工具",
    category=NodeCategory.DATA,
    icon="puzzle",
    description="调用 MCP (Model Context Protocol) 工具",
    inputs=[
        VariableDef(name="tool_name", type="string", required=True, description="工具名称"),
        VariableDef(name="arguments", type="object", description="工具参数"),
    ],
    outputs=[
        VariableDef(name="result", type="any", description="工具执行结果"),
        VariableDef(name="is_error", type="boolean", description="是否出错"),
    ],
    config_schema={
        "type": "object",
        "properties": {
            "server_url": {"type": "string", "description": "MCP 服务器地址"},
            "tool_name": {"type": "string", "description": "工具名称"},
            "arguments": {"type": "object", "description": "工具参数模板"},
            "timeout": {"type": "integer", "default": 30},
        },
        "required": ["server_url", "tool_name"],
    },
)
class MCPToolNode(BaseNode):
    """Calls MCP (Model Context Protocol) tools."""

    async def execute(self, variable_pool: Any, context: ExecutionContext | None = None) -> NodeResult:
        import httpx

        server_url = self.node_config.get("server_url", "")
        tool_name = self.node_config.get("tool_name", "")
        arguments_template = self.node_config.get("arguments", {})
        timeout = self.node_config.get("timeout", 30)

        # Resolve arguments templates
        arguments: dict[str, Any] = {}
        for key, value in arguments_template.items():
            if isinstance(value, str) and "{{" in value:
                arguments[key] = variable_pool.resolve_template(value)
            else:
                arguments[key] = value

        if not server_url or not tool_name:
            return NodeResult(status=NodeStatus.FAILED, error="Server URL and tool name are required")

        try:
            async with httpx.AsyncClient(timeout=timeout) as client:
                # MCP uses JSON-RPC 2.0
                response = await client.post(
                    server_url,
                    json={
                        "jsonrpc": "2.0",
                        "id": 1,
                        "method": "tools/call",
                        "params": {
                            "name": tool_name,
                            "arguments": arguments,
                        },
                    },
                )

                result_data = response.json()
                if "error" in result_data:
                    return NodeResult(
                        status=NodeStatus.FAILED,
                        error=result_data["error"].get("message", "MCP tool error"),
                    )

                tool_result = result_data.get("result", {})
                is_error = tool_result.get("isError", False)
                content = tool_result.get("content", [])

                return NodeResult(
                    status=NodeStatus.SUCCEEDED if not is_error else NodeStatus.FAILED,
                    outputs={"result": content, "is_error": is_error},
                )
        except Exception as e:
            logger.error("MCP tool execution failed", node_id=self.node_id, error=str(e))
            return NodeResult(status=NodeStatus.FAILED, error=str(e))


# ──────────────────────────────────────────────
# Wait Node
# ──────────────────────────────────────────────


@register_node(
    node_type="wait",
    display_name="等待",
    category=NodeCategory.DATA,
    icon="clock",
    description="暂停工作流执行，等待指定时间或外部信号",
    outputs=[
        VariableDef(name="waited_seconds", type="number", description="实际等待时间(秒)"),
    ],
    config_schema={
        "type": "object",
        "properties": {
            "wait_type": {
                "type": "string",
                "enum": ["duration", "signal", "timestamp"],
                "default": "duration",
            },
            "duration_seconds": {"type": "integer", "default": 0, "description": "等待时间(秒)"},
            "signal_name": {"type": "string", "description": "等待的信号名称"},
            "timestamp": {"type": "string", "description": "等待到指定时间"},
            "max_wait_seconds": {"type": "integer", "default": 3600, "description": "最大等待时间"},
        },
        "required": ["wait_type"],
    },
)
class WaitNode(BaseNode):
    """Pauses workflow execution for a duration, until a signal, or until a timestamp."""

    async def execute(self, variable_pool: Any, context: ExecutionContext | None = None) -> NodeResult:
        wait_type = self.node_config.get("wait_type", "duration")
        max_wait = self.node_config.get("max_wait_seconds", 3600)

        if wait_type == "duration":
            duration = self.node_config.get("duration_seconds", 0)
            if duration <= 0:
                return NodeResult(status=NodeStatus.SUCCEEDED, outputs={"waited_seconds": 0})
            duration = min(duration, max_wait)
            await asyncio.sleep(duration)
            return NodeResult(status=NodeStatus.SUCCEEDED, outputs={"waited_seconds": duration})

        elif wait_type == "signal":
            signal_name = self.node_config.get("signal_name", "continue")
            # Signal-based waiting is handled by the scheduler
            # The node sets a flag and the scheduler pauses until signal is received
            variable_pool.set(f"{self.node_id}.output.waiting_for_signal", signal_name)
            return NodeResult(
                status=NodeStatus.SUCCEEDED,
                outputs={"waited_seconds": 0, "signal": signal_name},
            )

        elif wait_type == "timestamp":
            from datetime import datetime, timezone
            timestamp_str = self.node_config.get("timestamp", "")
            if not timestamp_str:
                return NodeResult(status=NodeStatus.FAILED, error="No timestamp provided")
            try:
                target_time = datetime.fromisoformat(timestamp_str.replace("Z", "+00:00"))
                now = datetime.now(timezone.utc)
                wait_seconds = (target_time - now).total_seconds()
                if wait_seconds > 0:
                    wait_seconds = min(wait_seconds, max_wait)
                    await asyncio.sleep(wait_seconds)
                return NodeResult(status=NodeStatus.SUCCEEDED, outputs={"waited_seconds": max(0, wait_seconds)})
            except ValueError as e:
                return NodeResult(status=NodeStatus.FAILED, error=f"Invalid timestamp: {e}")

        return NodeResult(status=NodeStatus.FAILED, error=f"Unknown wait type: {wait_type}")
